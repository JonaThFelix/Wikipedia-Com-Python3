# Criado por Jonathan Felix
#Estudo de pesquisa no console e documentação do wikipédia em formato .docx

import wikipedia
import re
from docx import Document
import time

#Criação da Data para final do resumo
#from datetime import date
#data_atual = date.today()
#print(data_atual) comando mostra a data no code, mas não na impressão

name = input('Digite seu nome: ')
inst = input('Onde você estuda? ')
wikipedia.set_lang('pt')
tittle = input('Sobre o que você quer pesquisar: ? \n')
time.sleep(1)

print(f"Um minutinho {name} que estou achando sua pesquisa...")
time.sleep(2)
print("Lembrando que sua pesquisa estará salva automaticamente em formato Word em seu PC !")
time.sleep(1)
print("~~ coletando informações ~~")
time.sleep(1)
while True:
    try:
        wiki = wikipedia.page(tittle)
        break
    except:
        print(f"Estranho, não achei nenhuma referencia para: {tittle}")
        print("Tente novamente!!! ")
        tittle = input("Sobre o que você quer pesquisar? \n")

text = wiki.content
text = re.sub(r'==','',text)
text = re.sub(r'=','',text)
text = re.sub(r'\n','\n',text)
split = text.split('Veja também',1)
text = split[0]

print(text)

document = Document()
paragraph = document.add_heading(tittle,0)
paragraph.alignment = 2

paragraph = document.add_paragraph('' + text)
paragraph = document.add_paragraph(" ")
paragraph = document.add_paragraph("_________________________________________________________________________________________________________")
paragraph = document.add_paragraph("Pesquisa realizada por: " + name)
paragraph = document.add_paragraph("Instituição de ensino: " + inst)
paragraph = document.add_paragraph("Desenvolvido por: Jonathan Felix")
paragraph.aligment = 1
document.save(name + '-' + inst + '-'+ tittle + ".docx")
input()